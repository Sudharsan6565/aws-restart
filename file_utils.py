import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from io import BytesIO
from docx import Document as DocxDocument
from openpyxl import load_workbook
import csv

from langchain_core.documents import Document

def extract_text_from_pdf(path):
    texts = []
    try:
        doc = fitz.open(path)
        for i, page in enumerate(doc):
            page_text = page.get_text()
            for _, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image = Image.open(BytesIO(image_bytes)).convert("RGB")
                page_text += "\n" + pytesseract.image_to_string(image)
            if page_text.strip():
                texts.append(Document(page_content=page_text, metadata={"source": path, "page": i}))
    except Exception as e:
        print(f"[❌] Failed to parse PDF {path}: {e}")
    return texts

def extract_text_from_docx(path):
    try:
        doc = DocxDocument(path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return [Document(page_content=text, metadata={"source": path})] if text.strip() else []
    except Exception as e:
        print(f"[❌] Failed to parse DOCX {path}: {e}")
        return []

def extract_text_from_xlsx(path):
    try:
        wb = load_workbook(path, data_only=True)
        text = ""
        for sheet in wb.sheetnames:
            for row in wb[sheet].iter_rows(values_only=True):
                text += "\n" + " | ".join([str(cell) if cell else "" for cell in row])
        return [Document(page_content=text.strip(), metadata={"source": path})] if text.strip() else []
    except Exception as e:
        print(f"[❌] Failed to parse XLSX {path}: {e}")
        return []

def extract_text_from_csv(path):
    try:
        text = ""
        with open(path, newline="", encoding="utf-8") as csvfile:
            reader = csv.reader(csvfile)
            for row in reader:
                text += " | ".join(row) + "\n"
        return [Document(page_content=text.strip(), metadata={"source": path})] if text.strip() else []
    except Exception as e:
        print(f"[❌] Failed to parse CSV {path}: {e}")
        return []

def extract_text_from_image(path):
    try:
        image = Image.open(path).convert("RGB")
        text = pytesseract.image_to_string(image)
        print(f"[🧠 OCR TEXT] Extracted from {path}:\n{text}")
        return [Document(page_content=text.strip(), metadata={"source": path})] if text.strip() else []
    except Exception as e:
        print(f"[❌] Failed to parse image {path}: {e}")
        return []

def parse_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".xlsx":
        return extract_text_from_xlsx(file_path)
    elif ext == ".csv":
        return extract_text_from_csv(file_path)
    elif ext in [".txt", ".md"]:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
                return [Document(page_content=text.strip(), metadata={"source": file_path})] if text.strip() else []
        except Exception as e:
            print(f"[❌] Failed to read text file {file_path}: {e}")
            return []
    elif ext in [".png", ".jpg", ".jpeg", ".webp"]:
        return extract_text_from_image(file_path)

    return []
